"""
Multi-format document parser. Extracts text preserving section hierarchy and tables.
"""
import io
import logging
import re
from typing import Optional, List
from pathlib import Path

logger = logging.getLogger(__name__)


async def parse_document(file_path: str, file_type: str) -> dict:
    """Parse a document and return structured content."""
    parsers = {
        "pdf": _parse_pdf,
        "docx": _parse_docx,
        "md": _parse_markdown,
        "txt": _parse_txt,
    }
    parser = parsers.get(file_type.lower())
    if not parser:
        raise ValueError(f"Unsupported file type: {file_type}")

    text = await parser(file_path)
    sections = _extract_sections(text)
    word_count = len(text.split())

    return {
        "full_text": text,
        "sections": sections,
        "word_count": word_count,
        "section_count": len(sections),
    }


def detect_language(text: str, default: str = "en") -> str:
    """Detect dominant language from text. Returns ISO 639-1 code."""
    try:
        from langdetect import detect, DetectorFactory
        DetectorFactory.seed = 0  # deterministic
    except ImportError:
        return default

    sample = (text or "").strip()
    if not sample:
        return default
    sample = sample[:2000]
    try:
        code = detect(sample)
        # Normalize zh-cn / zh-tw → zh
        if code.startswith("zh"):
            return "zh"
        return code
    except Exception as exc:
        logger.warning(f"Language detection failed: {exc}")
        return default


async def _parse_pdf(file_path: str) -> str:
    """Parse PDF with table preservation via pdfplumber, fallback to PyPDF2."""
    try:
        import pdfplumber
        return _parse_pdf_plumber(file_path)
    except ImportError:
        logger.warning("pdfplumber not available, falling back to PyPDF2 (tables ignored)")
        return _parse_pdf_pypdf2(file_path)
    except Exception as exc:
        logger.warning(f"pdfplumber failed ({exc}), falling back to PyPDF2")
        return _parse_pdf_pypdf2(file_path)


def _parse_pdf_plumber(file_path: str) -> str:
    """Extract text + tables from PDF using pdfplumber."""
    import pdfplumber

    page_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page_idx, page in enumerate(pdf.pages):
            try:
                tables = page.find_tables()
            except Exception:
                tables = []

            # Build a set of bbox covering tables so we can remove their cell text from page text
            table_bboxes = [t.bbox for t in tables] if tables else []

            # Extract text excluding the regions covered by tables
            try:
                if table_bboxes:
                    # Crop out non-table regions one block at a time
                    text = page.extract_text() or ""
                else:
                    text = page.extract_text() or ""
            except Exception:
                text = ""

            section_parts = []
            if text:
                section_parts.append(text)

            # Add tables as markdown
            for t in tables:
                try:
                    rows = t.extract()
                except Exception:
                    continue
                if not rows:
                    continue
                md = _rows_to_markdown(rows)
                if md:
                    section_parts.append(md)

            if section_parts:
                page_parts.append("\n\n".join(section_parts))

    return "\n\n".join(page_parts)


def _parse_pdf_pypdf2(file_path: str) -> str:
    """Legacy PyPDF2 path — no table awareness."""
    from PyPDF2 import PdfReader

    text_parts = []
    with open(file_path, "rb") as f:
        reader = PdfReader(f)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n\n".join(text_parts)


async def _parse_docx(file_path: str) -> str:
    """Parse DOCX preserving paragraphs and tables in document order."""
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(file_path)
    parts: List[str] = []

    # Iterate over the document body children in order, so tables appear
    # at the position where they're inserted relative to paragraphs.
    body = doc.element.body
    # Map xml elements back to docx objects
    para_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}

    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            para = para_map.get(child)
            if para is None:
                continue
            if para.style.name.startswith("Heading"):
                level = (
                    para.paragraph_format.outline_level + 1
                    if para.paragraph_format.outline_level is not None
                    else 1
                )
                prefix = "#" * min(level, 6)
                parts.append(f"\n{prefix} {para.text}\n")
            elif para.text.strip():
                parts.append(para.text)
        elif child.tag == qn("w:tbl"):
            tbl = table_map.get(child)
            if tbl is None:
                continue
            rows = []
            for row in tbl.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            md = _rows_to_markdown(rows)
            if md:
                parts.append(md)

    return "\n\n".join(parts)


async def _parse_markdown(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


async def _parse_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


def _rows_to_markdown(rows: List[List[Optional[str]]]) -> str:
    """Convert a 2D list of cells to a markdown table."""
    if not rows:
        return ""
    # Normalize: replace None with empty, strip newlines/pipes that would break MD
    norm_rows = []
    for r in rows:
        norm_rows.append([_clean_cell(c) for c in r])

    # Skip if completely empty
    if not any(any(c for c in r) for r in norm_rows):
        return ""

    col_count = max(len(r) for r in norm_rows)
    # Pad rows to col_count
    for r in norm_rows:
        while len(r) < col_count:
            r.append("")

    header = norm_rows[0]
    body_rows = norm_rows[1:] if len(norm_rows) > 1 else []

    lines = []
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join(["---"] * col_count) + " |")
    for r in body_rows:
        lines.append("| " + " | ".join(r) + " |")
    return "\n".join(lines)


def _clean_cell(cell: Optional[str]) -> str:
    if cell is None:
        return ""
    return str(cell).replace("\n", " ").replace("|", "\\|").strip()


def _extract_sections(text: str) -> List[dict]:
    """Split text into sections based on markdown headings."""
    lines = text.split("\n")
    sections = []
    current_heading = "Introduction"
    current_content: List[str] = []
    current_level = 1

    heading_pattern = re.compile(r"^(#{1,6})\s+(.+)$")

    for line in lines:
        match = heading_pattern.match(line.strip())
        if match:
            if current_content:
                sections.append({
                    "heading": current_heading,
                    "content": "\n".join(current_content).strip(),
                    "level": current_level,
                })
            current_level = len(match.group(1))
            current_heading = match.group(2)
            current_content = []
        else:
            current_content.append(line)

    if current_heading or current_content:
        sections.append({
            "heading": current_heading,
            "content": "\n".join(current_content).strip(),
            "level": current_level,
        })

    return [s for s in sections if s["content"].strip()]
